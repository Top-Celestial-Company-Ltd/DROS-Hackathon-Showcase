import docx

doc = docx.Document(r'E:\vscode\AI知識庫\DROS-Hackathon-Showcase\temp_team.docx')
lines = []

for p in doc.paragraphs:
    if p.text.strip():
        lines.append(p.text.strip())

for t in doc.tables:
    lines.append("\n--- TABLE ---")
    for row in t.rows:
        r_text = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
        lines.append(' | '.join(r_text))

with open(r'E:\vscode\AI知識庫\DROS-Hackathon-Showcase\temp_extracted.txt', 'w', encoding='utf-8') as f:
    f.write('\n'.join(lines))
print("EXTRACTED_SUCCESSFULLY")
